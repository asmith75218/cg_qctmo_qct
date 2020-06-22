# OOI/CGSN - WHOI
# ---------------
# cg_ctdmo_qct.py
# version 2.1
# 22 June 2020
# ---------------------------------------------------------------------------- #
#
# Automates the QCT procedure for OOI CTDMO instruments (SBE 37IM), and outputs
# calibration CSV file for asset management.
#
# Written for python 3.x
# with pyserial and python-docx libraries
#
# Please see README.md for installation and usage information.
#
# 

import re
import csv
import serial
from serial.tools import list_ports
from datetime import datetime as dt
from datetime import timedelta
from docx import Document

import cal_ctdmo

# ---- Function definitions ----
def set_formnumber(formnumber):
    """Prompt the user for a form number, or increment it if it already exists."""
    if not formnumber:
        # Prompt user and test for valid input...
        input_validated = False
        while not input_validated:
            formnumber = input("Enter the five-digit form number (ex: 00123): ")
            if re.match("^\d{5}$", formnumber) and formnumber != "00000":
                input_validated = True
            else:
                print("Error! Invalid entry.\r\n")
        
        return formnumber
    else:
        # increment the form number...
        return str(int(formnumber)+1).rjust(5, '0')

def select_port():
    """
    Display a list of available ports and prompt the user to
    select which to use. If only one port is available, that one
    will be used automatically, with no prompt.
    """
    while True:
        ports = list_ports.comports()
        if len(ports) == 0:
            print("There don't seem to be any available serial ports.")
            selection = input("Would you like to connect a serial device and try again? [y]/n ")
            if selection != "n":
                input("Connect your device now, then press ENTER to continue...")
            else:
                return None
        else:
            break

    if len(ports) == 1:
        return ports[0][0]
    while True:
        print("Available ports:")
        for n, port in enumerate(ports, 1):
            print("%d) %s" % (n, port))
        # Input testing -- input has to be a number.
        while True:
            try:
                selection = int(input("Enter your selection: "))
            except ValueError:
                print("Whoa! That's not a number.")
                continue
            break
        # Input testing -- input has to be in list_ports range.
        try:
            return ports[selection-1][0]
        except IndexError:
            print("Whoa! That's not an available port.")
    

def ser_coms(ser, capfile, cmd):
    """
    Function to transmit a message over a serial connection,
    capture the received input and write the input to the capfile.

    Args:
        ser: an open pySerial connection
        capfile: an open file
        cmd (str): message to transmit

    Returns:
        str: the received message.
    """
    ser.reset_input_buffer()
    ser.write((cmd + '\r\n').encode('ascii'))
    cap = ser.read(ser.in_waiting).decode('ascii')
    while True:
        prev = cap
        cap += ser.read(1).decode('ascii')
        if prev == cap:
            capfile.write(cap)
            return cap

def open_port(port, baudrate):
    """Function to open a serial connection."""
    print("Connecting to %s at %d baud..." % (port, baudrate))
    try:
        ser = serial.Serial(port, baudrate, timeout=5)
        print("Connected to %s." % port)
        return ser
    except:
        return None

def get_remote_id(ser, capfile):
    """
    Establish communication with the instrument and get its id no.

    Return the two-digit id no. if successful.
    Return "None" if communication fails.
    """
    while True:
        id_query = ser_coms(ser, capfile, 'id?')
        # Search for id using regex, not really necessary when an id is returned
        # but useful for when it fails to get an id back from the instrument, so
        # I've left it for now...
        remote_id = re.search('id\s=\s(\d{2})', id_query)
        if not remote_id:
            print("Unable to get id from remote instrument.")
            if input("Would you like to try again? [y]/n ") == "n":
                return None
        else:
            return remote_id.group(1)

def reset_remote_id(ser, capfile, desired_id):
    """
    Function to reset the instrument's id no. to a new value.

    Return the two-digit new id if successful.
    Return "None" if communication fails.
    """
    while True:
        response = ser_coms(ser, capfile, '*id=%s' % desired_id)
        if 'FAILED' in response:
            print("Unable to communicate with remote instrument.")
            if input("Would you like to try again? [y]/n ") == "n":
                return None
        elif 'id = %s' % desired_id in response:
            return desired_id

def split_csv(csv):
    """
    Returns values from a sample as a list of strings.
    """
    # When an instrument takes a sample, it returns one line of
    # comma separated data, along with some other characters on
    # other lines. This code extracts the second line, which is
    # most likely the data, splits it at commas and removes
    # whitespace padding. THIS WILL NOT WORK ON MULTIPLE LINES
    # OF DATA. It will probably just take the first one.
    return [x.strip() for x in csv.splitlines()[1].split(',')]

def date_from_ds(ds):
    """Extract the date from the DS message and return as datetime obj."""
    ds_date = ' '.join(ds.split()[6:10])
    return dt.strptime(ds_date, '%d %b %Y %H:%M:%S')
    
def get_from_ts(ts):
    """Strip out non-useful characters and return a list with just the data"""
    # TODO(gsmith) Test here for a complete line so the formatting works
    return [x.strip() for x in ts.replace('\n', '').replace('\r', '')[5:-13].split(',')]

def tidy_up(ser, capfile):
    """
    Close the serial port and all open files when test is aborted or complete.
    """
    ser_coms(ser, capfile, 'pwroff')    # Put the instruments to sleep
    capfile.close()    # Close the capture file
    ser.close()    # Release the serial port
    return
    
def dict_from_csv(csvfilename):
    """
    Import and convert a csv file to a python dictionary. Each line in the csv should
    be in the format "key, value".
    """
    with open(csvfilename, 'r') as csvfile:
        return dict(csv.reader(csvfile))
    
# ---- Test loop ----
def ctdmo_qct_test(port, username, formnumber):
    """
    QCT test procedure.

    Args:
        port: An available serial port.
        username (str): Name of person conducting the test.
        formnumber (str): Five-digit number for the results document.

    Returns:
        str: A status message.
    """
    # Function variables
    csvfilename = 'ctdmo_inv.csv'
    partnumber_dict = {'G': '00007',  'H': '00008',  'Q': '00017',  'R': '00018'}
    
    # ---- Test Setup ----
    # Open serial port...
    #port = select_port()
    ser = open_port(port, 9600)
    if not ser:
        return "Error! Unable to open port!"

    # Open capture file...
    capfile = open(("3305-00101-%s-A.txt" % formnumber), 'w')

    # ---- 8.3.5 ----
    print("Waking the IMM...")
    ser_coms(ser, capfile, 'pwron')

    # ---- 8.3.6 ----
    print("Establishing communication with the remote device...")
    remote_id = get_remote_id(ser, capfile)
    if remote_id == None:
        tidy_up(ser, capfile)
        return "Attention! Test cancelled by user."
    else:
        print("Remote id: %s" % remote_id)

    while remote_id != desired_id:
        print("Resetting remote id...")
        remote_id = reset_remote_id(ser, capfile, desired_id)
        if remote_id == None:
            tidy_up(ser, capfile)
            return "Error! Test cancelled by user."
        print("Remote id: %s" % remote_id)

    # ---- Set up results document ----
    # Make a list for test results at each step...
    step_results_p = [False for _ in range(10)]

    # Make a default list for 'Test Data' column entries...
    testdata = [None for _ in range(10)]

    # The first step is already complete...
    step_results_p[0] = True
    testdata[0] = "Established communication with instrument."

    # Open the results document (MS Word)...
    doc = Document('3305-00101-00000.docx')

    # Calling doc.tables will return all tables in the template, where all
    # text will go. Access each of these by index number. Start with the
    # first one, at index 0...
    table = doc.tables[0]

    # ---- 8.3.7 ----
    ser_coms(ser, capfile, '#%soutputformat=1' % remote_id)
    ds = ser_coms(ser, capfile, '#%sds' % remote_id).split()
    serialnumber = "37-%s" % ds[5]
    firmware = ds[2]
    print("Serial number: %s" % serialnumber)
    print("Firmware version: %s" % firmware)

    step_results_p[1] = True
    testdata[1] = "Serial number confirmed."
    step_results_p[2] = True
    testdata[2] = "Firmware %s confirmed." % firmware

    # Use serial number to look up series letter from a csv inventory file...
    inv_dict = dict_from_csv(csvfilename)
    try:
        seriesletter = inv_dict[serialnumber]
    except KeyError:
        while True:
            seriesletter = input("Enter the instrument Series (G, H, Q or R): ").upper()
            if seriesletter in "GHQR": break
            print("Whoa! That's not a valid series letter,  %s." % username.split()[0])
    print("Class/Series: CTDMO-%s" % seriesletter)
    print("Part number: 1336-00001-%s" % partnumber_dict[seriesletter])
    
    # ---- Results document header ----
    # A little out of order, but with the serial number extracted
    # the document header can now be completed.
    
    # partnumber and serialnumber are found in the third row, first cell...
    cell = table.rows[2].cells[0]

    # Since each line is technically a separate paragraph, fill them in
    # one at a time to better preserve formatting...
    cell.paragraphs[1].text = cell.paragraphs[1].text.replace('partnumber', partnumber_dict[seriesletter])
    cell.paragraphs[3].text = cell.paragraphs[3].text.replace('serialnumber', serialnumber)

    # seriesletter is found in the third row, second cell...
    cell = table.rows[2].cells[1]
    cell.paragraphs[1].text = cell.paragraphs[1].text.replace('seriesletter', seriesletter)

    # formnumber is found in the third row, last cell...
    cell = table.rows[2].cells[-1]
    cell.paragraphs[2].text = cell.paragraphs[2].text.replace('formnumber', formnumber)

    # On to the next table...
    table = doc.tables[1]
    cell = table.rows[0].cells[0]

    # Fill in the username...
    cell.paragraphs[1].text = cell.paragraphs[1].text.replace('username', username)

    # Fill in the date...
    cell.paragraphs[1].text = cell.paragraphs[1].text.replace('testdate', dt.today().strftime("%Y-%m-%d"))

    # On to the test steps and results table...
    table = doc.tables[2]

    # ---- 8.3.8 ----
    print("Retrieving calibration information...")
    cc_repl = ser_coms(ser, capfile, '#%sgetcc' % remote_id)

    # Generate the calibration CSV...
    print("Exporting calibration to CSV...")
    cc_xml = cc_repl[8:-2]
    cal_ctdmo.generate_csv(cc_xml, seriesletter, formnumber)
    
    # ---- 8.3.9 ----
    this_step_complete = False
    while not this_step_complete:
        print("Testing instrument clock...")
        testdata[3] = "Test not completed"
        rollback_target = (dt.utcnow() - timedelta(days=1)).replace(hour=12,
                                                             minute=0, second=0)
        ser_coms(ser, capfile, '#%sdatetime=%s'
                 % (remote_id, rollback_target.strftime('%m%d%Y%H%M%S')))
        rollback_instr = date_from_ds(ser_coms(ser, capfile, '#%sds' % remote_id))
        # Test if clock was changed...
        # (Some time will have passed between when we set the instrument
        # clock to noon, and when we queried the instrument again. We don't
        # really know what this is, but with two 5-sec timeouts, we have
        # set this margin at 15 sec. If instruments are repeatedly failing
        # and the clock is confirmed to be set correctly, then it may be
        # necessary to change the 15 to something higher on the next line)
        if not 0 <= (rollback_instr - rollback_target).seconds <= 15:
            print("The instrument time %s is not set to noon yesterday."
                  % rollback_instr.strftime('%d %b %Y %H:%M:%S'))
            response = input("Does this bother you, %s? [y]/n "
                                 % username.split()[0])
            if response != "n":
                response = input("Would you like to try again? [y]/n ")
                if response == "n":
                    testdata[3] = ("The instrument clock was not set "
                                   "successfully.")
                else: continue
             

    # ---- 8.3.10 ----
        ser_coms(ser, capfile, '#%sdatetime=%s'
                 % (remote_id, (dt.utcnow().strftime('%m%d%Y%H%M%S'))))
        currenttime_instr = date_from_ds(ser_coms(ser, capfile, '#%sds' % remote_id))
        # Test if clock was reset correctly...
        # (As above, there will be a slight discrepency resulting from the
        # time between commands. It is set here at 15 sec. Change on the next
        # line if deemed necessary)
        if not 0 <= (dt.utcnow() - currenttime_instr).seconds <= 15:
            print("The instrument time %s is not set to the current time %s."
                  % (currenttime_instr.strftime('%d %b %Y %H:%M:%S'),
                     dt.utcnow().strftime('%d %b %Y %H:%M:%S')))
            response = input("Does this bother you, %s? [y]/n "
                                 % username.split()[0])
            if response != "n":
                response = input("Would you like to try again? [y]/n ")
                if response == "n":
                    if testdata[3]:
                       testdata[3] = testdata[3] + ("\r\nThe instrument clock was "
                                                    "not reset successfully.")
                    else:
                        testdata[3] = ("The instrument clock was not reset "
                                   "successfully.")
                    this_step_complete = True
            else:
                step_results_p[3] = True
                testdata[3] = "Instrument clock set successfully."
                this_step_complete = True     
        else:
            step_results_p[3] = True
            testdata[3] = "Instrument clock set successfully."
            this_step_complete = True     
            
 
    # ---- 8.3.11 ----
    this_step_complete = False
    while not this_step_complete:
        print("Testing configuration...")
        testdata[4] = "Test not completed."
        desired_sample_interval = "120"
        ser_coms(ser, capfile, '#%ssampleinterval=%s'
                 % (remote_id, desired_sample_interval))
        ds = ser_coms(ser, capfile, '#%sds' % remote_id).split()
        sample_interval = ds[29]
        print(' '.join(ds[26:31]))
        
        # Update test results...
        if sample_interval != desired_sample_interval:
            print("The sample interval was not set to the correct value.")
            response = input("Does this bother you, %s? [y]/n "
                             % username.split()[0])
            if response != "n":
                response = input("Would you like to try again? [y]/n ")
                if response == "n":
                    testdata[4] = ("The sample interval was not set "
                                   "successfully.\r\n")
                    this_step_complete = True
            else:
                step_results_p[4] = True
                testdata[4] = ("Sampling configuration changed successfully."
                               "\r\n%s" % ' '.join(ds[26:31]))
                this_step_complete = True     
        else:
            step_results_p[4] = True
            testdata[4] = ("Sampling configuration changed successfully."
                           "\r\n%s" % ' '.join(ds[26:31]))
            this_step_complete = True     

    # ---- 8.3.12 ----
    this_step_complete = False
    while not this_step_complete:
        print("Configuring instrument...")
        if step_results_p[5]: step_results_p[5] = False
        testdata[5] = "Test not completed."
        desired_sample_interval = "10"
        ser_coms(ser, capfile, '#%ssampleinterval=%s'
                 % (remote_id, desired_sample_interval))
        ds = ser_coms(ser, capfile, '#%sds' % remote_id).split()
        sample_interval = ds[29]
        print(' '.join(ds[26:31]))

        # Update test results...
        if sample_interval != desired_sample_interval:
            print("The sample interval was not set to the correct value.")
            response = input("Does this bother you, %s? [y]/n "
                             % username.split()[0])
            if response != "n":
                response = input("Would you like to try again? [y]/n ")
                if response == "n":
                    testdata[5] = ("The sample interval was not set "
                                   "successfully.\r\n")
                    this_step_complete = True
            else:
                step_results_p[5] = True
                testdata[5] = ("Sampling configuration changed successfully."
                               "\r\n%s" % ' '.join(ds[26:31]))
                this_step_complete = True     
        else:
            step_results_p[5] = True
            testdata[5] = ("Sampling configuration changed successfully."
                           "\r\n%s" % ' '.join(ds[26:31]))
            this_step_complete = True     
    
    # ---- 8.3.13 ----
    this_step_complete = False
    while not this_step_complete:
        print("Acquiring sample...")
        step_results_p[6:8] = [False] * 3
        testdata[6:8] = ["Test not completed."] * 3
        sample_air = split_csv(ser_coms(ser, capfile, '#%sts' % remote_id))

        # Test for date stamp...
        if dt.strptime(sample_air[4], '%d %b %Y').date() != dt.utcnow().date():
            print("The sample date of %s does not appear to be correct."
                  % (' '.join(sample_air[4:6])))
            response = input("Does this bother you, %s? [y]/n "
                                 % username.split()[0])
            if response != "n":
                response = input("Would you like to try again? [y]/n ")
                if response == "n":    # Update test results...
                    testdata[6] = ("The sample does not contain a valid "
                                   "date stamp.\r\n%s" % ', '.join(sample_air))
                else: continue
            else:
                step_results_p[6] = True
                testdata[6] = "Correct date stamp confirmed."
        else:
            step_results_p[6] = True
            testdata[6] = "Correct date stamp confirmed."

        # Test for reasonable values...
        if (float(sample_air[1]) < 10 or    # Lower bound room temp
            float(sample_air[1]) > 30 or    # Upper bound room temp
            float(sample_air[2]) < -1 or    # Lower bound cond
            float(sample_air[2]) > 1 or     # Upper bound cond
            float(sample_air[3]) < -2 or    # Lower bound atm press
            float(sample_air[3]) > 2):      # Upper bound atm press

            print("The sample data don't seem quite right.\r\n%s"
                  % (' '.join(sample_air)))
            response = input("Does this bother you, %s? [y]/n "
                                 % username.split()[0])
            if response != "n":
                response = input("Would you like to try again? [y]/n ")
                if response == "n":    # Update test results...
                    testdata[7] = ("The sample data are invalid.\r\n%s"
                                   % ', '.join(sample_air))
                else: continue
            else:
                step_results_p[7] = True
                testdata[7] = ("Sample values seem reasonable.\r\n%s"
                               % ', '.join(sample_air))
        else:
            step_results_p[7] = True
            testdata[7] = ("Sample values seem reasonable.\r\n%s"
                           % ', '.join(sample_air))
        
    # ---- 8.3.15 ----
    # ---- 8.3.16 ----
        print("Place the instrument in a container of warm water now.")
        input("Press ENTER to continue...")
        print("Acquiring sample...")
        sample_bucket = split_csv(ser_coms(ser, capfile, '#%sts' % remote_id))

        # Compare the results...
        # Did the pressure change as expected?
        if float(sample_bucket[3]) <= float(sample_air[3]):
            print("The bucket pressure of %s is not greater than the air pressure "
                  "of %s as expected!" % (sample_bucket[3], sample_air[3]))
            response = input("Does this bother you, %s? [y]/n "
                                 % username.split()[0])
            if response != "n":
                response = input("Would you like to try again? [y]/n ")
                if response == "n":    # Update test results...
                    testdata[8] = ("The pressure did not increase as expected.\r\n%s"
                                   % ', '.join(sample_bucket))
                else:
                    print("Remove the instrument from the water now.")
                    input("Press ENTER to continue...")
                    continue
            else:
                step_results_p[8] = True
                testdata[8] = "Pressure sensor correctly records changes."
        else:
            step_results_p[8] = True
            testdata[8] = "Pressure sensor correctly records changes."

        # Did the temperature change as expected?
        if float(sample_bucket[1]) - float(sample_air[1]) < 1.0:
            print("The bucket temperature of %sC is not warmer than the air "
                  "temperature of %sC as expected!" % (sample_bucket[1],
                                                       sample_air[1]))
            response = input("Does this bother you, %s? [y]/n "
                                 % username.split()[0])
            if response != "n":
                response = input("Would you like to try again? [y]/n ")
                if response == "n":    # Update test results...
                    testdata[9] = ("The temperature did not increase as expected.\r\n%s"
                                   % ', '.join(sample_bucket))
                    this_step_complete = True
            else:
                step_results_p[9] = True
                testdata[9] = "Temp data changes as expected."
                this_step_complete = True     
        else:
            step_results_p[9] = True
            testdata[9] = "Temp data changes as expected."
            this_step_complete = True     

    # ---- Wrap up the results document ----
    # Populate the Pass and Fail columns...
    for i in range(10):
        if step_results_p[i]:
            # True case: PASS
            table.columns[4].cells[i+2].text = testdata[i]
            table.columns[5].cells[i+2].text = 'X'
            table.columns[6].cells[i+2].text = ''
        else:
            # False case: FAIL
            table.columns[4].cells[i+2].text = testdata[i]
            table.columns[5].cells[i+2].text = ''
            table.columns[6].cells[i+2].text = 'X'

    # Complete the last table...
    table = doc.tables[4]
    table.rows[1].cells[0].text = ("Session log file: 3305-00101-%s-A.txt"
                                   % formnumber)

    # Set the document Title and Author attributes...
    doc.core_properties.title = ("SN_%s_QCT_Results_CTDMO-%s"
                                 % (serialnumber, seriesletter))
    author_inits = username.split()
    author_inits[:-1] = [init[0] + '.' for init in author_inits[:-1]]
    doc.core_properties.author = ' '.join(author_inits)

    # Save a copy of the results file with the form number in the title...
    doc.save('3305-00101-%s.docx' % formnumber)

    # ---- Test Complete ----
    tidy_up(ser, capfile)
    return "Test complete."

# ---- Variable definitions ----
time_to_quit = False
port = None
formnumber = None
desired_id = '01'

# ---------------------------------------------------------------------------- #
# ---- Main program loop ----
while not time_to_quit:
    print("\r\nCTDMO QCT v2.0.1")
    print("MAIN MENU")
    print("---------")
    print("1) Test an instrument")
    print("2) Configure serial port")
    print("3) Exit")
    selection = input("Enter your selection: ")
    if selection == '1':
        if not port:
            port = select_port()
            if not port:
                continue
        username = input("What is your name?: ")
        while True:
            formnumber = set_formnumber(formnumber)
            print(ctdmo_qct_test(port, username, formnumber))
            again = input("Would you like to test another instrument? y/[n] ")
            if again != "y":
                formnumber = None
                break
            input("Type ENTER to begin the next test...")
    elif selection == '2':
        port = select_port()
        print("Selected port is %s." % port)
    elif selection == '3':
        time_to_quit = True
        print("Good bye!")
    else: print("Error! Invalid entry.\r\n")
